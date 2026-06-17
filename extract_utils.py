"""Text extraction helpers for uploaded literature files."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from pypdf import PdfReader


class UnsupportedFileTypeError(ValueError):
    """Raised when the uploaded file type is not supported."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(BytesIO(file_bytes))
    page_texts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            page_texts.append(f"[第 {index} 页]\n{text.strip()}")
    return "\n\n".join(page_texts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract paragraph and table text from a Word .docx file."""
    document = Document(BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text_from_upload(file_name: str, file_bytes: bytes) -> str:
    """Extract text from a supported uploaded file."""
    lower_name = file_name.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    raise UnsupportedFileTypeError("当前仅支持 PDF 和 Word .docx 文件。")
